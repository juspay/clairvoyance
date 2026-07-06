"""
File-upload connector: GCS object -> NormalizedContent.

Parsing is format-aware: PDF (pypdf), DOCX (python-docx, including embedded
tables), Markdown/plain text (passthrough), CSV (stdlib) and XLSX (openpyxl)
as tabular data. Parsers run in a thread (they're CPU-bound / the GCS client
is sync) so the event loop stays free.
"""

import asyncio
import csv
import hashlib
import io
from typing import List, Tuple

from app.core.logger import logger
from app.schemas.breeze_buddy.knowledge_base import KbDocument
from app.services.gcp.storage.storage import GCSStorage
from app.services.knowledge_base.connectors.base import KBConnector
from app.services.knowledge_base.types import NormalizedContent, TableData

MAX_TABLE_COLUMNS = 50


def _extension(name: str) -> str:
    """Lowercased file extension ('' if none) — used by ``_parse`` as the
    fallback when the mime type is missing or too generic to dispatch on."""
    return name.rsplit(".", 1)[-1].lower() if "." in name else ""


def _parse_pdf(data: bytes) -> NormalizedContent:
    """PDF bytes -> plain text (page texts joined by blank lines).

    A page that fails extraction is skipped with a warning rather than
    failing the document — scanned/image pages yield empty strings, and
    ``load`` rejects the document only if NOTHING was extractable. Import
    is deferred so pypdf loads only when a PDF is actually ingested.
    """
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:
            logger.warning(f"pypdf failed to extract a page: {e}")
    return NormalizedContent(text="\n\n".join(pages))


def _parse_docx(data: bytes) -> NormalizedContent:
    """DOCX bytes -> markdown-ish text + embedded tables.

    Word heading styles become ``#`` markdown headings so the prose chunker
    can split on sections and build breadcrumbs; embedded tables (>= 2 rows)
    are lifted out as TableData for row-as-chunk treatment.
    """
    import docx

    document = docx.Document(io.BytesIO(data))
    # Preserve heading structure as markdown so the prose chunker can split
    # on sections.
    lines: List[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            try:
                level = int(style.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            lines.append("#" * min(level, 6) + " " + text)
        else:
            lines.append(text)

    tables: List[TableData] = []
    for table_index, table in enumerate(document.tables):
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) >= 2:
            tables.append(
                TableData(
                    name=f"table_{table_index + 1}",
                    headers=rows[0][:MAX_TABLE_COLUMNS],
                    rows=[r[:MAX_TABLE_COLUMNS] for r in rows[1:]],
                )
            )
    return NormalizedContent(text="\n".join(lines), tables=tables)


def _parse_csv(data: bytes, name: str) -> NormalizedContent:
    """CSV bytes -> one TableData (first row = headers).

    utf-8-sig handles Excel's BOM exports; blank rows are dropped; the
    filename (sans extension) becomes the table name that row chunks are
    prefixed with.
    """
    text = data.decode("utf-8-sig", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = [row for row in reader if any(cell.strip() for cell in row)]
    if not rows:
        return NormalizedContent()
    table_name = name.rsplit(".", 1)[0]
    return NormalizedContent(
        tables=[
            TableData(
                name=table_name,
                headers=rows[0][:MAX_TABLE_COLUMNS],
                rows=[r[:MAX_TABLE_COLUMNS] for r in rows[1:]],
            )
        ]
    )


def _parse_xlsx(data: bytes) -> NormalizedContent:
    """XLSX bytes -> one TableData per worksheet (first row = headers).

    ``read_only`` + ``data_only`` keeps memory flat on big workbooks and
    reads computed values instead of formulas. Sheets with fewer than 2
    rows (no data under the header) are skipped.
    """
    import openpyxl

    workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    tables: List[TableData] = []
    try:
        for sheet in workbook.worksheets:
            rows: List[List[str]] = []
            for row in sheet.iter_rows(values_only=True):
                cells = ["" if v is None else str(v) for v in row]
                if any(cell.strip() for cell in cells):
                    rows.append(cells[:MAX_TABLE_COLUMNS])
            if len(rows) >= 2:
                tables.append(
                    TableData(name=sheet.title, headers=rows[0], rows=rows[1:])
                )
    finally:
        workbook.close()
    return NormalizedContent(tables=tables)


def _parse_text(data: bytes) -> NormalizedContent:
    """TXT/Markdown bytes -> text passthrough (markdown headings are
    understood downstream by the prose chunker's section splitter)."""
    return NormalizedContent(text=data.decode("utf-8", errors="replace"))


def _parse(data: bytes, name: str, mime_type: str) -> NormalizedContent:
    """Dispatch raw bytes to the right format parser.

    Mime type wins when informative, extension is the fallback (browsers
    send unreliable mime types for CSVs especially). Raises ValueError for
    unsupported types — belt-and-braces behind the API's extension check.
    """
    ext = _extension(name)
    mime = (mime_type or "").lower()

    if "pdf" in mime or ext == "pdf":
        return _parse_pdf(data)
    if "wordprocessingml" in mime or ext == "docx":
        return _parse_docx(data)
    if "spreadsheetml" in mime or ext == "xlsx":
        return _parse_xlsx(data)
    if "csv" in mime or ext == "csv":
        return _parse_csv(data, name)
    if ext in ("txt", "md", "markdown") or mime.startswith("text/"):
        return _parse_text(data)

    raise ValueError(
        f"Unsupported file type for '{name}' (mime: {mime_type or 'unknown'}). "
        "Supported: pdf, docx, xlsx, csv, txt, md"
    )


class FileUploadConnector(KBConnector):
    source_type = "file"

    async def load(self, document: KbDocument) -> NormalizedContent:
        """Download the uploaded file from GCS and parse it.

        The connector half of ingestion's ``_process_document``. Download +
        parse run in a worker thread (sync GCS client, CPU-bound parsers) so
        the event loop — shared with live calls — stays free. Sets
        ``content_hash`` from the raw bytes for whole-document change checks.
        """
        gcs_path = str(document.source_ref.get("gcs_path") or "")
        if not gcs_path:
            raise ValueError(f"Document {document.id} has no gcs_path in source_ref")

        def _download_and_parse() -> Tuple[NormalizedContent, str]:
            # Runs inside asyncio.to_thread: both the GCS download (sync
            # client) and parsing (CPU-bound) happen off the event loop.
            data = GCSStorage().download_as_bytes(gcs_path)
            if data is None:
                raise RuntimeError(f"Failed to download {gcs_path} from GCS")
            parsed = _parse(data, document.name, document.mime_type or "")
            return parsed, hashlib.sha256(data).hexdigest()

        content, content_hash = await asyncio.to_thread(_download_and_parse)
        content.content_hash = content_hash

        if not content.text and not content.tables:
            raise ValueError(
                f"No extractable text found in '{document.name}' -- the file may "
                "be scanned/image-only or empty"
            )
        return content

    async def detect_change(self, document: KbDocument) -> bool:
        """Uploaded files are immutable snapshots; a new upload creates a new
        GCS object. Manual re-sync re-processes unconditionally."""
        return True
